import base64
import io
import json
import math
import os
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from PIL import Image
from wordcloud import WordCloud

from src.file_readers.file_readers_docx import read_docx
from src.file_readers.file_readers_pdf import read_pdf
from src.file_readers.file_readers_txt import read_txt
from src.skill_extractor import extract_skills
from src.text_cleaner.remove_personal import remove_personal
from src.text_cleaner.section_normalizer import normalize_text as full_normalize

try:
    from askill_ext import AdvancedSkillExtractor
except Exception:
    AdvancedSkillExtractor = None

try:
    from gap_analysys import (
        GapAnalysisResult,
        GapVisualizer,
        LearningPathGenerator,
        SentenceBERTEncoder,
        SimilarityCalculator,
        SkillGapAnalyzer,
    )
    MILESTONE3_AVAILABLE = True
except Exception:
    GapAnalysisResult = None
    GapVisualizer = None
    LearningPathGenerator = None
    SentenceBERTEncoder = None
    SimilarityCalculator = None
    SkillGapAnalyzer = None
    MILESTONE3_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False


st.set_page_config(
    page_title="AI Skill Gap Analyzer Pro",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)


APP_DIR = Path(__file__).resolve().parent
SKILLS_FILE = APP_DIR / "skills_list.txt"
DEMO_RESUME_CANDIDATES = [
    APP_DIR / "mohan_resume.docx",
    APP_DIR / "mohan_resume.pdf",
    APP_DIR / "mohan_resume.txt",
]
DEMO_JD_CANDIDATES = [
    APP_DIR / "mohan_job_description.pdf",
    APP_DIR / "mohan_job_description.docx",
    APP_DIR / "mohan_job_description.txt",
]


TECHNICAL_KEYWORDS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "c",
    "c++",
    "c#",
    "sql",
    "html",
    "css",
    "react",
    "node",
    "node.js",
    "django",
    "flask",
    "fastapi",
    "machine learning",
    "deep learning",
    "artificial intelligence",
    "ai",
    "nlp",
    "bert",
    "spacy",
    "pytorch",
    "tensorflow",
    "scikit-learn",
    "sklearn",
    "pandas",
    "numpy",
    "matplotlib",
    "seaborn",
    "plotly",
    "mysql",
    "postgresql",
    "mongodb",
    "git",
    "github",
    "docker",
    "kubernetes",
    "aws",
    "gcp",
    "azure",
    "tableau",
    "power bi",
    "data science",
    "data analysis",
    "data analytics",
    "etl",
    "api",
    "rest",
    "problem solving",
}
SOFT_KEYWORDS = {
    "leadership",
    "communication",
    "collaboration",
    "teamwork",
    "adaptability",
    "creativity",
    "time management",
    "critical thinking",
    "analytical skills",
    "project management",
    "decision making",
    "mentoring",
    "negotiation",
    "organization",
    "planning",
    "innovation",
}
TOOL_KEYWORDS = {
    "git",
    "github",
    "vscode",
    "visual studio code",
    "eclipse",
    "power bi",
    "tableau",
    "excel",
    "jupyter",
    "docker",
    "kubernetes",
    "jira",
    "confluence",
    "linux",
    "bash",
    "shell",
    "aws",
    "azure",
    "gcp",
    "mysql",
    "postgresql",
    "mongodb",
}
CERTIFICATION_KEYWORDS = {
    "certification",
    "certified",
    "certificate",
    "nptel",
    "coursera",
    "udemy",
    "edx",
    "forage",
    "freecodecamp",
    "azure ai fundamentals",
}

DISPLAY_OVERRIDES = {
    "js": "JavaScript",
    "ts": "TypeScript",
    "sql": "SQL",
    "ai": "AI",
    "ml": "ML",
    "nlp": "NLP",
    "aws": "AWS",
    "gcp": "GCP",
    "github": "GitHub",
    "power bi": "Power BI",
    "node.js": "Node.js",
    "scikit-learn": "Scikit-learn",
}

CATEGORY_ORDER = ["technical", "soft", "tools", "certifications"]
CATEGORY_LABELS = {
    "technical": "Technical",
    "soft": "Soft",
    "tools": "Tools",
    "certifications": "Certifications",
    "other": "Other",
}
CATEGORY_COLORS = {
    "technical": "#2E86FF",
    "soft": "#FF8B1A",
    "tools": "#35B44B",
    "certifications": "#B7B7B7",
    "other": "#7C8795",
}


@dataclass
class DocumentResult:
    name: str
    doc_type: str
    raw_text: str
    cleaned_text: str
    extracted_skills: List[str]
    advanced_skills: List[str]
    skill_confidence: Dict[str, float]
    original_chars: int
    cleaned_chars: int
    removed_lines: int
    removed_chars: int
    word_count: int
    file_size_kb: float
    pages: int
    source_kind: str


@dataclass
class AnalysisBundle:
    resume: DocumentResult
    job_description: DocumentResult
    all_resume_skills: List[str]
    all_jd_skills: List[str]
    analysis_result: Optional[object]
    match_score: float
    ats_score: float
    ats_factors: Dict[str, Dict[str, object]]
    matched_skills: List[Dict[str, object]]
    partial_matches: List[Dict[str, object]]
    missing_skills: List[Dict[str, object]]
    learning_path: List[Dict[str, object]]
    category_counts_resume: Dict[str, int]
    category_counts_jd: Dict[str, int]
    similarity_matrix: Optional[np.ndarray]
    skill_pair_map: List[Dict[str, object]]



def load_custom_css() -> None:
    st.markdown(
        """
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@400;600;700;800&family=Space+Grotesk:wght@500;700&display=swap');

:root {
    --bg: #0b0f17;
    --panel: #171d2a;
    --panel-2: #20283a;
    --line: rgba(175, 200, 255, 0.16);
    --text: #eef3ff;
    --muted: rgba(218,229,255,0.72);
    --accent: #6f7ef8;
    --accent-2: #8a63de;
    --good: #2ecc71;
    --warn: #ffb020;
    --bad: #ff4d7d;
}

.stApp,
body,
[data-testid="stAppViewContainer"] {
    font-family: "Sora", "Segoe UI", sans-serif;
}

.main {
    background:
        radial-gradient(70% 60% at 0% -5%, rgba(95, 141, 255, 0.24), transparent 55%),
        radial-gradient(60% 50% at 100% 0%, rgba(129, 98, 219, 0.22), transparent 52%),
        linear-gradient(180deg, #070b12 0%, #0a0f19 100%);
}

.block-container {
    padding-top: 1rem;
    padding-bottom: 2rem;
    max-width: 1500px;
}

header[data-testid="stHeader"] {
    background: rgba(14,16,20,0.92);
}

section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1e222b 0%, #191d24 100%);
    border-right: 1px solid rgba(255,255,255,0.06);
}

.hero {
    position: relative;
    border-radius: 22px;
    overflow: hidden;
    min-height: 140px;
    padding: 28px 34px;
    background: linear-gradient(120deg, rgba(105,112,232,0.96), rgba(120,96,214,0.92) 45%, rgba(113,147,255,0.86));
    box-shadow: 0 20px 45px rgba(5, 9, 18, 0.55);
    border: 1px solid rgba(219, 230, 255, 0.2);
    display: grid;
    place-items: center;
    text-align: center;
}

.hero:after {
    content: "";
    position: absolute;
    inset: 0;
    background: linear-gradient(128deg, rgba(255,255,255,0.18), rgba(255,255,255,0.01) 45%, rgba(20,32,66,0.2) 75%);
    pointer-events: none;
}

.hero h1 {
    margin: 0 0 0.35rem;
    color: white;
    font-size: clamp(2.05rem, 3.6vw, 3.2rem);
    line-height: 1.04;
    font-weight: 800;
    letter-spacing: -0.03em;
    text-shadow: 0 4px 14px rgba(0,0,0,0.32);
}

.hero p {
    margin: 0;
    color: rgba(255,255,255,0.92);
    font-size: 0.95rem;
    font-weight: 500;
}

.hero-badge {
    display: inline-flex;
    align-items: center;
    gap: 0.4rem;
    border: 1px solid rgba(255,255,255,0.25);
    background: rgba(13, 18, 39, 0.2);
    color: rgba(255,255,255,0.94);
    border-radius: 999px;
    padding: 0.36rem 0.72rem;
    margin-bottom: 0.75rem;
    font-size: 0.78rem;
    font-weight: 700;
}

.panel {
    background: rgba(30,34,42,0.92);
    border: 1px solid var(--line);
    border-radius: 18px;
    padding: 1rem 1.15rem;
    box-shadow: 0 10px 30px rgba(0,0,0,0.22);
}

.section-title {
    color: var(--text);
    font-size: 1.75rem;
    font-weight: 800;
    margin: 0.3rem 0 0.8rem;
    letter-spacing: -0.025em;
    font-family: "Space Grotesk", "Sora", sans-serif;
}

.section-subtitle {
    color: var(--muted);
    margin-bottom: 1rem;
}

.stepper {
    display: grid;
    grid-template-columns: repeat(6, 1fr);
    gap: 8px;
    margin: 0;
    position: relative;
    z-index: 2;
}

.step-shell {
    position: relative;
    padding: 0.7rem 0.45rem 0.5rem;
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 16px;
    background: linear-gradient(180deg, rgba(32,38,52,0.7), rgba(24,30,43,0.78));
}

.step-track {
    position: absolute;
    left: 6%;
    right: 6%;
    top: 35px;
    height: 5px;
    border-radius: 999px;
    background: linear-gradient(90deg, rgba(100,120,255,0.45), rgba(95,95,120,0.2));
    border: 1px solid rgba(255,255,255,0.08);
}

.step-card {
    text-align: center;
    color: var(--text);
}

.step-circle {
    width: 48px;
    height: 48px;
    margin: 0 auto 8px;
    border-radius: 999px;
    display: grid;
    place-items: center;
    font-weight: 700;
    font-size: 1rem;
    background: #242d3f;
    border: 1px solid rgba(220,230,255,0.24);
    color: rgba(240,246,255,0.8);
}

.step-circle.active {
    background: linear-gradient(180deg, rgba(131,120,233,1), rgba(104,88,210,1));
    color: white;
    box-shadow: 0 0 0 4px rgba(127,108,242,0.18), 0 0 18px rgba(127,108,242,0.55);
}

.step-circle.done {
    background: linear-gradient(180deg, #31c46d, #1fa458);
    color: white;
}

.step-label {
    font-size: 0.78rem;
    color: rgba(255,255,255,0.90);
    font-weight: 700;
}

.step-icon {
    display: block;
    margin-top: 3px;
    font-size: 0.72rem;
    color: rgba(255,255,255,0.62);
}

.nav-buttons {
    display: grid;
    grid-template-columns: repeat(6, 1fr);
    gap: 10px;
    margin: 0.7rem 0 1rem;
}

.nav-buttons button {
    width: 100%;
    border-radius: 10px;
    border: 1px solid rgba(255,255,255,0.08);
    background: linear-gradient(180deg, rgba(112,126,245,0.92), rgba(111,90,214,0.88));
    color: white;
    padding: 0.42rem 0.5rem;
    font-size: 0.72rem;
}

.upload-card {
    background: rgba(25,29,36,0.92);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 18px;
    padding: 0.9rem 1rem 1rem;
}

.upload-zone {
    border: 1.5px dashed rgba(97,127,255,0.58);
    border-radius: 16px;
    min-height: 96px;
    background: rgba(255,255,255,0.03);
    padding: 0.85rem;
}

.card-title {
    font-size: 1.02rem;
    font-weight: 700;
    color: var(--text);
    margin-bottom: 0.55rem;
}

.small-note {
    color: var(--muted);
    font-size: 0.86rem;
}

.status-card {
    background: linear-gradient(180deg, rgba(41,48,62,0.84), rgba(34,40,53,0.78));
    border: 1px solid rgba(196,214,255,0.14);
    border-radius: 18px;
    padding: 1rem 1.05rem;
    margin: 0.7rem 0 1rem;
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.02), 0 10px 24px rgba(0,0,0,0.25);
}

.status-card.good {
    border-left: 4px solid #31c46d;
}

.status-card.warn {
    border-left: 4px solid #ffb020;
}

.metric-grid {
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    gap: 12px;
}

.metric-box {
    background: linear-gradient(180deg, rgba(20,50,87,0.76), rgba(18,41,69,0.72));
    border: 1px solid rgba(121,167,255,0.24);
    border-radius: 18px;
    padding: 0.95rem 1rem;
    text-align: center;
    box-shadow: 0 12px 24px rgba(0,0,0,0.22);
}

.metric-big {
    color: white;
    font-size: 2.1rem;
    font-weight: 800;
    margin-bottom: 0.4rem;
}

.metric-label {
    color: rgba(255,255,255,0.85);
    font-size: 0.93rem;
    font-weight: 600;
}

.chip-wrap {
    display: flex;
    flex-wrap: wrap;
    gap: 10px;
    margin-top: 0.7rem;
}

.chip {
    display: inline-flex;
    align-items: center;
    gap: 0.35rem;
    padding: 0.42rem 0.88rem;
    border-radius: 999px;
    font-size: 0.82rem;
    font-weight: 700;
    color: white;
    box-shadow: 0 8px 18px rgba(0,0,0,0.15);
}

.chip.good { background: linear-gradient(180deg, #31c46d, #1fa458); }
.chip.warn { background: linear-gradient(180deg, #ffbf36, #f28b1f); }
.chip.bad { background: linear-gradient(180deg, #ff5683, #e62f5f); }
.chip.neutral { background: linear-gradient(180deg, #6a77ff, #4d62d8); }

.side-panel {
    background: rgba(255,255,255,0.03);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 16px;
    padding: 0.9rem;
    margin-bottom: 1rem;
}

.side-panel h3 {
    color: white;
    font-size: 0.96rem;
    margin-bottom: 0.7rem;
}

.side-panel label,
.side-panel p,
.side-panel div {
    color: rgba(255,255,255,0.86);
}

.factors-grid {
    display: grid;
    grid-template-columns: repeat(2, minmax(0, 1fr));
    gap: 14px;
}

.ats-shell {
    background: linear-gradient(180deg, rgba(26,32,45,0.96), rgba(21,27,40,0.96));
    border: 1px solid rgba(147,171,218,0.14);
    border-radius: 18px;
    padding: 1.2rem 1.1rem;
}

.ats-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 1rem;
}

.ats-title {
    color: white;
    font-size: 1.02rem;
    font-weight: 700;
}

.ats-bar {
    width: 100%;
    height: 9px;
    background: rgba(255,255,255,0.08);
    border-radius: 999px;
    overflow: hidden;
    margin: 0.55rem 0 0.2rem;
}

.ats-bar > div {
    height: 100%;
    border-radius: inherit;
}

.ats-recommendation {
    margin-top: 0.5rem;
    color: rgba(255,255,255,0.72);
    font-size: 0.85rem;
}

.score-card {
    background: linear-gradient(135deg, rgba(45,49,58,0.95), rgba(26,29,36,0.95));
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 20px;
    padding: 2.1rem 1.5rem;
    text-align: center;
    position: relative;
    overflow: hidden;
}

.score-card:before {
    content: "";
    position: absolute;
    inset: 0;
    background: linear-gradient(135deg, rgba(255,255,255,0.08), transparent 40%);
    pointer-events: none;
}

.score-circle {
    width: 112px;
    height: 112px;
    border-radius: 999px;
    display: grid;
    place-items: center;
    margin: 0 auto 1rem;
    font-weight: 900;
    font-size: 2rem;
    color: white;
    border: 4px solid rgba(255,255,255,0.22);
    box-shadow: 0 0 0 8px rgba(255,255,255,0.03), 0 0 26px rgba(255,169,34,0.28);
}

.score-circle.excellent { background: linear-gradient(180deg, #31c46d, #1fa458); }
.score-circle.good { background: linear-gradient(180deg, #5bb6ff, #366fd8); }
.score-circle.average { background: linear-gradient(180deg, #ffbf36, #f28b1f); }
.score-circle.poor { background: linear-gradient(180deg, #ff5683, #e62f5f); }

.report-shell {
    background: #ffffff;
    border-radius: 10px;
    padding: 1.2rem 1.3rem;
    color: #1d2430;
    box-shadow: 0 16px 36px rgba(0,0,0,0.24);
}

.report-shell h1,
.report-shell h2,
.report-shell h3,
.report-shell h4 {
    color: #26334a;
}

.report-shell table {
    width: 100%;
    border-collapse: collapse;
}

.report-shell th,
.report-shell td {
    border: 1px solid #d7ddea;
    padding: 0.55rem 0.6rem;
    text-align: left;
}

.report-shell th {
    background: #eef3ff;
}

.footer-note {
    color: rgba(255,255,255,0.6);
    font-size: 0.83rem;
    text-align: center;
    margin-top: 1.5rem;
}

@media (max-width: 1200px) {
    .metric-grid,
    .factors-grid {
        grid-template-columns: repeat(2, minmax(0, 1fr));
    }

    .stepper,
    .nav-buttons {
        grid-template-columns: repeat(3, 1fr);
    }
}

@media (max-width: 768px) {
    .hero {
        padding: 20px 16px;
        min-height: 132px;
    }

    .metric-grid,
    .factors-grid,
    .stepper,
    .nav-buttons {
        grid-template-columns: 1fr;
    }
}
</style>
        """,
        unsafe_allow_html=True,
    )


@st.cache_data(show_spinner=False)
def read_skill_list() -> List[str]:
    if not SKILLS_FILE.exists():
        return []
    try:
        return [line.strip() for line in SKILLS_FILE.read_text(encoding="utf-8").splitlines() if line.strip()]
    except Exception:
        return []


@st.cache_resource(show_spinner=False)
def get_advanced_extractor_instance():
    if AdvancedSkillExtractor is None:
        return None
    try:
        return AdvancedSkillExtractor()
    except Exception:
        return None


@st.cache_resource(show_spinner=False)
def get_sentence_encoder(model_name: str):
    if SentenceBERTEncoder is None:
        return None
    try:
        return SentenceBERTEncoder(model_name)
    except Exception:
        return None


@st.cache_resource(show_spinner=False)
def get_similarity_calculator():
    if SimilarityCalculator is None:
        return None
    try:
        return SimilarityCalculator()
    except Exception:
        return None


@st.cache_resource(show_spinner=False)
def get_learning_generator():
    if LearningPathGenerator is None:
        return None
    try:
        return LearningPathGenerator()
    except Exception:
        return None


@st.cache_resource(show_spinner=False)
def get_gap_visualizer():
    return GapVisualizer if GapVisualizer is not None else None



def reset_analysis() -> None:
    keys = [
        "bundle",
        "processed_docs",
        "resume_text_input",
        "jd_text_input",
        "resume_upload_name",
        "jd_upload_name",
        "resume_file",
        "jd_file",
        "nav_target",
        "loading_demo",
        "demo_loaded",
        "last_report_bytes",
        "last_report_name",
    ]
    for key in keys:
        if key in st.session_state:
            del st.session_state[key]



def ensure_state() -> None:
    defaults = {
        "nav_target": "upload",
        "analysis_mode": "spaCy base model",
        "confidence_threshold": 0.60,
        "similarity_threshold": 0.60,
        "embedding_model": "all-MiniLM-L6-v2",
        "export_format": "PDF",
        "demo_loaded": False,
        "last_report_bytes": None,
        "last_report_name": None,
        "bundle": None,
        "processed_docs": [],
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)



def clean_skill_name(skill: str) -> str:
    token = re.sub(r"\s+", " ", str(skill)).strip()
    if not token:
        return token
    lowered = token.lower()
    if lowered in DISPLAY_OVERRIDES:
        return DISPLAY_OVERRIDES[lowered]
    if token.isupper() and len(token) <= 5:
        return token
    if "+" in token or "#" in token or "." in token:
        return token
    return " ".join(part.capitalize() if part.isalpha() else part for part in token.split())



def canonical_skill_key(skill: str) -> str:
    token = re.sub(r"\s+", " ", str(skill)).strip().lower()
    return DISPLAY_OVERRIDES.get(token, token).lower()



def unique_preserve_order(values: List[str]) -> List[str]:
    seen = set()
    ordered = []
    for value in values:
        if not value:
            continue
        key = canonical_skill_key(value)
        if key not in seen:
            seen.add(key)
            ordered.append(clean_skill_name(value))
    return ordered



def extract_text_from_bytes(file_name: str, file_bytes: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    try:
        if suffix == ".txt":
            try:
                return file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                return file_bytes.decode("latin-1", errors="ignore")
        if suffix == ".docx":
            import docx

            document = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()])
        if suffix == ".pdf":
            try:
                import pdfplumber

                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    pages = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages).strip()
                if text:
                    return text
            except Exception:
                pass
            try:
                import PyPDF2

                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                pages = [page.extract_text() or "" for page in reader.pages]
                return "\n".join(pages).strip()
            except Exception:
                return ""
    except Exception:
        return ""
    return ""



def estimate_pages(file_name: str, raw_text: str, file_bytes: bytes) -> int:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".pdf":
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return max(1, len(pdf.pages))
        except Exception:
            pass
    if suffix == ".docx":
        try:
            import docx

            document = docx.Document(io.BytesIO(file_bytes))
            paragraphs = len([p.text for p in document.paragraphs if p.text.strip()])
            return max(1, math.ceil(max(1, paragraphs) / 12))
        except Exception:
            pass
    return max(1, math.ceil(max(1, len(raw_text.split())) / 300))



def count_removed_lines(original: str, cleaned: str) -> int:
    original_lines = [line for line in original.splitlines() if line.strip()]
    cleaned_lines = [line for line in cleaned.splitlines() if line.strip()]
    return max(0, len(original_lines) - len(cleaned_lines))



def get_demo_file() -> Tuple[Optional[Path], Optional[Path]]:
    resume_path = next((path for path in DEMO_RESUME_CANDIDATES if path.exists()), None)
    jd_path = next((path for path in DEMO_JD_CANDIDATES if path.exists()), None)
    return resume_path, jd_path



def get_sample_inputs_from_workspace() -> Optional[Tuple[Dict[str, object], Dict[str, object]]]:
    resume_path, jd_path = get_demo_file()
    if not resume_path or not jd_path:
        return None
    try:
        resume_bytes = resume_path.read_bytes()
        jd_bytes = jd_path.read_bytes()
        return (
            {"name": resume_path.name, "bytes": resume_bytes, "kind": "resume"},
            {"name": jd_path.name, "bytes": jd_bytes, "kind": "job_description"},
        )
    except Exception:
        return None



def parse_and_clean_document(name: str, file_bytes: bytes, kind: str) -> DocumentResult:
    raw_text = extract_text_from_bytes(name, file_bytes)
    if not raw_text.strip():
        raise ValueError(f"Could not extract text from {name}")

    cleaned_text = full_normalize(remove_personal(raw_text))
    if not cleaned_text.strip():
        raise ValueError(f"Cleaning produced empty text for {name}")

    original_chars = len(raw_text)
    cleaned_chars = len(cleaned_text)
    removed_chars = max(0, original_chars - cleaned_chars)
    removed_lines = count_removed_lines(raw_text, cleaned_text)
    file_size_kb = round(len(file_bytes) / 1024, 1)
    pages = estimate_pages(name, raw_text, file_bytes)
    word_count = len(cleaned_text.split())

    return DocumentResult(
        name=name,
        doc_type=kind,
        raw_text=raw_text,
        cleaned_text=cleaned_text,
        extracted_skills=[],
        advanced_skills=[],
        skill_confidence={},
        original_chars=original_chars,
        cleaned_chars=cleaned_chars,
        removed_lines=removed_lines,
        removed_chars=removed_chars,
        word_count=word_count,
        file_size_kb=file_size_kb,
        pages=pages,
        source_kind=kind,
    )



def extract_skills_hybrid(cleaned_text: str, extractor_mode: str) -> Tuple[List[str], List[str], Dict[str, float]]:
    base_skills = []
    advanced_skills = []
    confidence_map: Dict[str, float] = {}

    if SKILLS_FILE.exists():
        try:
            base_skills = extract_skills(cleaned_text, str(SKILLS_FILE))
        except Exception:
            base_skills = []

    if extractor_mode == "AI hybrid model":
        extractor = get_advanced_extractor_instance()
        if extractor is not None:
            try:
                advanced_result = extractor.extract_skills(cleaned_text, document_type="resume")
                if advanced_result.get("success"):
                    advanced_skills = advanced_result.get("all_skills", [])
                    confidence_map = advanced_result.get("skill_confidence", {})
            except Exception:
                advanced_skills = []
                confidence_map = {}

    merged = unique_preserve_order(base_skills + advanced_skills)
    if not merged and cleaned_text.strip():
        # final fallback for basic keyword capture if the skill list is sparse
        text_lower = cleaned_text.lower()
        fallback_hits = []
        for skill in read_skill_list():
            if skill.lower() in text_lower:
                fallback_hits.append(skill)
        merged = unique_preserve_order(fallback_hits)

    return merged, unique_preserve_order(advanced_skills), confidence_map



def categorize_skill(skill: str) -> str:
    token = skill.lower().strip()
    if any(keyword == token or keyword in token for keyword in CERTIFICATION_KEYWORDS):
        return "certifications"
    if any(keyword == token or keyword in token for keyword in SOFT_KEYWORDS):
        return "soft"
    if any(keyword == token or keyword in token for keyword in TOOL_KEYWORDS):
        return "tools"
    if any(keyword == token or keyword in token for keyword in TECHNICAL_KEYWORDS):
        return "technical"
    return "other"



def build_category_counts(skills: List[str]) -> Dict[str, int]:
    counts = defaultdict(int)
    for skill in skills:
        counts[categorize_skill(skill)] += 1
    return {category: counts.get(category, 0) for category in CATEGORY_ORDER}



def create_tag_html(skills: List[str], category: str) -> str:
    if not skills:
        return '<div class="small-note">No skills detected.</div>'

    color_class = {
        "technical": "good",
        "soft": "warn",
        "tools": "neutral",
        "certifications": "bad",
        "other": "neutral",
    }.get(category, "neutral")

    chips = []
    for skill in skills:
        chips.append(f'<span class="chip {color_class}">{skill}</span>')
    return '<div class="chip-wrap">' + "".join(chips) + "</div>"



def similarity_for_pair(a: str, b: str) -> float:
    a_clean = re.sub(r"[^a-z0-9]+", "", a.lower())
    b_clean = re.sub(r"[^a-z0-9]+", "", b.lower())
    if not a_clean or not b_clean:
        return 0.0
    if a_clean == b_clean:
        return 1.0
    return SequenceMatcher(None, a_clean, b_clean).ratio()



def build_fallback_similarity(resume_skills: List[str], jd_skills: List[str], threshold: float) -> Tuple[List[Dict[str, object]], List[Dict[str, object]], List[Dict[str, object]], np.ndarray, float]:
    matrix = np.zeros((max(1, len(resume_skills)), max(1, len(jd_skills))))
    matched, partial, missing = [], [], []
    for j_idx, jd_skill in enumerate(jd_skills):
        best_idx = 0
        best_score = -1.0
        for r_idx, resume_skill in enumerate(resume_skills):
            score = similarity_for_pair(resume_skill, jd_skill)
            matrix[r_idx, j_idx] = score
            if score > best_score:
                best_idx = r_idx
                best_score = score
        best_resume = resume_skills[best_idx] if resume_skills else ""
        bucket = {
            "jd_skill": jd_skill,
            "resume_skill": best_resume,
            "similarity": float(best_score if best_score >= 0 else 0),
            "category": "STRONG_MATCH" if best_score >= 0.8 else ("PARTIAL_MATCH" if best_score >= threshold else "MISSING"),
            "confidence_level": "HIGH" if best_score >= 0.8 else ("MEDIUM" if best_score >= threshold else "LOW"),
            "priority": "LOW" if best_score >= 0.8 else ("MEDIUM" if best_score >= threshold else "HIGH"),
        }
        if best_score >= 0.8:
            matched.append(bucket)
        elif best_score >= threshold:
            partial.append(bucket)
        else:
            missing.append(bucket)
    overall_score = ((len(matched) + (0.5 * len(partial))) / len(jd_skills)) if jd_skills else 0.0
    return matched, partial, missing, matrix, overall_score



def run_semantic_gap_analysis(resume_skills: List[str], jd_skills: List[str], threshold: float):
    if not resume_skills or not jd_skills:
        return None, [], [], [], np.zeros((0, 0)), 0.0

    if MILESTONE3_AVAILABLE:
        encoder = get_sentence_encoder("all-MiniLM-L6-v2")
        calculator = get_similarity_calculator()
        if encoder is not None and calculator is not None:
            try:
                analyzer = SkillGapAnalyzer(
                    encoder,
                    calculator,
                    strong_threshold=0.80,
                    partial_threshold=threshold,
                )
                categories = {skill: categorize_skill(skill) for skill in jd_skills}
                result = analyzer.analyze(resume_skills, jd_skills, skill_categories=categories)
                matched = [
                    {
                        "jd_skill": item.jd_skill,
                        "resume_skill": item.resume_skill,
                        "similarity": float(item.similarity),
                        "category": item.category,
                        "confidence_level": item.confidence_level,
                        "priority": item.priority,
                    }
                    for item in result.matched_skills
                ]
                partial = [
                    {
                        "jd_skill": item.jd_skill,
                        "resume_skill": item.resume_skill,
                        "similarity": float(item.similarity),
                        "category": item.category,
                        "confidence_level": item.confidence_level,
                        "priority": item.priority,
                    }
                    for item in result.partial_matches
                ]
                missing = [
                    {
                        "jd_skill": item.jd_skill,
                        "resume_skill": item.resume_skill,
                        "similarity": float(item.similarity),
                        "category": item.category,
                        "confidence_level": item.confidence_level,
                        "priority": item.priority,
                    }
                    for item in result.missing_skills
                ]
                score = float(result.overall_score)
                return result, matched, partial, missing, result.similarity_matrix, score
            except Exception:
                pass

    matched, partial, missing, matrix, score = build_fallback_similarity(resume_skills, jd_skills, threshold)
    return None, matched, partial, missing, matrix, score



def build_learning_path(missing_skills: List[Dict[str, object]]) -> List[Dict[str, object]]:
    if not missing_skills:
        return []

    ranking = sorted(missing_skills, key=lambda item: (item.get("priority", "HIGH"), -item.get("similarity", 0)))
    path = []
    for item in ranking[:5]:
        skill = item["jd_skill"]
        priority = item.get("priority", "MEDIUM").title()
        importance = max(1.0, round((1 / max(1, len(ranking))) * 100, 1))
        skill_lower = skill.lower()
        if any(token in skill_lower for token in ["python", "java", "javascript", "react", "sql", "aws", "docker", "ml", "ai"]):
            estimated_time = "2-4 weeks" if priority != "High" else "4-8 weeks"
        else:
            estimated_time = "1-3 weeks"
        resources = [
            f"Online course: {skill} for Beginners",
            f"Practice projects on {skill}",
            f"Certification: {skill} Professional",
        ]
        path.append(
            {
                "skill": skill,
                "priority": priority,
                "importance": f"{importance}%",
                "estimated_time": estimated_time,
                "recommended_action": f"Learn {skill} through online courses and practical experience",
                "resources": resources,
            }
        )
    return path



def calculate_ats_factors(resume_text: str, jd_text: str, resume_skills: List[str], jd_skills: List[str], bundle: AnalysisBundle) -> Tuple[float, Dict[str, Dict[str, object]]]:
    resume_tokens = set(re.findall(r"[a-zA-Z0-9+.#-]+", resume_text.lower()))
    jd_tokens = set(re.findall(r"[a-zA-Z0-9+.#-]+", jd_text.lower()))

    resume_words = set(word.lower() for word in resume_text.split())
    jd_words = set(word.lower() for word in jd_text.split())

    keyword_overlap = len(set(skill.lower() for skill in jd_skills) & set(skill.lower() for skill in resume_skills))
    keyword_score = round(100 * keyword_overlap / max(1, len(jd_skills)), 1)

    sections = ["summary", "experience", "education", "projects", "skills", "certifications"]
    section_hits = sum(1 for section in sections if section in resume_text.lower())
    section_score = round((section_hits / len(sections)) * 100, 1)

    formatting_penalty = 0
    if len(re.findall(r"\n\s*[-•*]", resume_text)) >= 3:
        formatting_penalty += 5
    if len(re.findall(r"[|]{2,}", resume_text)) > 0:
        formatting_penalty += 5
    if len(re.findall(r"[A-Z]{4,}", resume_text)) > 20:
        formatting_penalty += 10
    formatting_score = max(50, 100 - formatting_penalty)

    avg_sentence_len = np.mean([len(sentence.split()) for sentence in re.split(r"[.!?]+", resume_text) if sentence.strip()]) if resume_text.strip() else 0
    readability_score = max(0, min(100, 100 - abs(avg_sentence_len - 16) * 3))

    experience_score = 100 if any(token in resume_words for token in jd_tokens) else 72
    if any(skill.lower() in resume_text.lower() for skill in jd_skills[:8]):
        experience_score = min(100, experience_score + 10)

    ats_score = round(
        keyword_score * 0.40
        + formatting_score * 0.20
        + section_score * 0.20
        + readability_score * 0.10
        + experience_score * 0.10,
        1,
    )

    factor_scores = {
        "keyword_match": {
            "score": keyword_score,
            "category": "excellent" if keyword_score >= 75 else "average" if keyword_score >= 50 else "poor",
            "recommendation": "Add exact job keywords naturally in your summary, projects, and experience sections.",
        },
        "formatting": {
            "score": formatting_score,
            "category": "excellent" if formatting_score >= 80 else "average" if formatting_score >= 60 else "poor",
            "recommendation": "Keep the resume readable: simple headings, clean bullets, and consistent spacing.",
        },
        "section_completeness": {
            "score": section_score,
            "category": "excellent" if section_score >= 80 else "average" if section_score >= 60 else "poor",
            "recommendation": "Include all key ATS sections such as Summary, Skills, Experience, Education, Projects, and Certifications.",
        },
        "readability": {
            "score": round(readability_score, 1),
            "category": "excellent" if readability_score >= 80 else "average" if readability_score >= 60 else "poor",
            "recommendation": "Use short, direct sentences and highlight measurable outcomes.",
        },
        "experience_relevance": {
            "score": round(experience_score, 1),
            "category": "excellent" if experience_score >= 80 else "average" if experience_score >= 60 else "poor",
            "recommendation": "Mirror the job description with project bullets that show direct relevance.",
        },
    }

    return ats_score, factor_scores



def score_badge_class(score: float) -> str:
    if score >= 80:
        return "excellent"
    if score >= 60:
        return "good"
    if score >= 40:
        return "average"
    return "poor"



def generate_similarity_heatmap(resume_skills: List[str], jd_skills: List[str], similarity_matrix: Optional[np.ndarray]):
    if similarity_matrix is None or similarity_matrix.size == 0:
        fig = go.Figure()
        fig.update_layout(
            height=460,
            template="plotly_dark",
            paper_bgcolor="#111318",
            plot_bgcolor="#111318",
            annotations=[dict(text="No similarity data available", x=0.5, y=0.5, xref="paper", yref="paper", showarrow=False)],
        )
        return fig

    max_display = min(20, len(resume_skills), len(jd_skills))
    matrix = similarity_matrix[:max_display, :max_display]
    fig = go.Figure(
        data=go.Heatmap(
            z=matrix,
            x=jd_skills[:max_display],
            y=resume_skills[:max_display],
            colorscale="RdYlGn",
            zmin=0,
            zmax=1,
            colorbar=dict(title="Similarity"),
            text=np.round(matrix, 2),
            texttemplate="%{text}",
            hovertemplate="Resume: %{y}<br>JD: %{x}<br>Similarity: %{z:.2f}<extra></extra>",
        )
    )
    fig.update_layout(
        height=620,
        template="plotly_dark",
        paper_bgcolor="#111318",
        plot_bgcolor="#111318",
        margin=dict(l=30, r=20, t=30, b=30),
    )
    return fig



def generate_radar_chart(category_resume: Dict[str, int], category_jd: Dict[str, int]):
    labels = [CATEGORY_LABELS[c] for c in CATEGORY_ORDER]
    resume_values = [category_resume.get(c, 0) for c in CATEGORY_ORDER]
    jd_values = [category_jd.get(c, 0) for c in CATEGORY_ORDER]

    if sum(resume_values) == 0:
        resume_values = [0, 0, 0, 0]
    if sum(jd_values) == 0:
        jd_values = [0, 0, 0, 0]

    fig = go.Figure()
    fig.add_trace(
        go.Scatterpolar(
            r=resume_values + [resume_values[0]],
            theta=labels + [labels[0]],
            fill="toself",
            name="Resume",
            line=dict(color="#3a66ff", width=2),
            fillcolor="rgba(58,102,255,0.30)",
        )
    )
    fig.add_trace(
        go.Scatterpolar(
            r=jd_values + [jd_values[0]],
            theta=labels + [labels[0]],
            fill="toself",
            name="Job Description",
            line=dict(color="#ff4d5f", width=2),
            fillcolor="rgba(255,77,95,0.20)",
        )
    )
    fig.update_layout(
        polar=dict(bgcolor="#111318", radialaxis=dict(visible=True, color="rgba(255,255,255,0.25)"), angularaxis=dict(color="rgba(255,255,255,0.5)")),
        paper_bgcolor="#111318",
        plot_bgcolor="#111318",
        template="plotly_dark",
        showlegend=True,
        height=520,
        margin=dict(l=20, r=20, t=30, b=20),
    )
    return fig



def generate_distribution_chart(category_resume: Dict[str, int]):
    labels = [CATEGORY_LABELS[c] for c in CATEGORY_ORDER]
    values = [category_resume.get(c, 0) for c in CATEGORY_ORDER]
    colors = [CATEGORY_COLORS[c] for c in CATEGORY_ORDER]
    fig = go.Figure(
        data=[
            go.Bar(
                x=labels,
                y=values,
                marker=dict(color=colors),
                text=values,
                textposition="auto",
            )
        ]
    )
    fig.update_layout(
        title="Skill Distribution by Category",
        paper_bgcolor="#111318",
        plot_bgcolor="#111318",
        template="plotly_dark",
        height=480,
        margin=dict(l=20, r=20, t=50, b=20),
        xaxis_title="Category",
        yaxis_title="Count",
    )
    return fig



def generate_wordcloud_image(skills: List[str]) -> Optional[np.ndarray]:
    if not skills:
        return None
    frequencies = Counter(skill if len(skill) > 1 else skill.upper() for skill in skills)
    if not frequencies:
        return None
    wordcloud = WordCloud(width=1600, height=860, background_color="white", colormap="viridis", prefer_horizontal=0.9).generate_from_frequencies(frequencies)
    return wordcloud.to_array()



def generate_pdf_report_bytes(bundle: AnalysisBundle) -> Optional[bytes]:
    if not REPORTLAB_AVAILABLE:
        return None

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=30)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleAccent", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20, textColor=colors.HexColor("#5f7df7"), spaceAfter=10)
    subtitle_style = ParagraphStyle("SubtitleAccent", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, textColor=colors.HexColor("#222222"), spaceAfter=10)
    section_style = ParagraphStyle("SectionAccent", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14, textColor=colors.HexColor("#2c2c2c"), spaceAfter=8)
    body_style = ParagraphStyle("BodyAccent", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12, textColor=colors.HexColor("#333333"))

    story = [
        Paragraph("Skill Gap Analysis Report", title_style),
        Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style),
        Paragraph("Overall Match Score", section_style),
        Table(
            [["Match Percentage", f"{bundle.match_score:.1f}%"], ["ATS Score", f"{bundle.ats_score:.1f}%"]],
            colWidths=[2.2 * inch, 1.4 * inch],
        ),
    ]
    story.append(Spacer(1, 10))

    skills_table = Table(
        [["Category", "Resume Skills", "JD Skills", "Matched"], ["Technical", len(bundle.all_resume_skills), len(bundle.all_jd_skills), len(bundle.matched_skills)]],
        colWidths=[1.2 * inch, 1.2 * inch, 1.0 * inch, 0.9 * inch],
    )
    skills_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef3ff")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cfd8ea")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story.append(Paragraph("Skills Summary", section_style))
    story.append(skills_table)
    story.append(Spacer(1, 10))

    def section_list(title: str, items: List[Dict[str, object]], color: str):
        story.append(Paragraph(title, section_style))
        if not items:
            story.append(Paragraph("No items found.", body_style))
            story.append(Spacer(1, 8))
            return
        for item in items[:8]:
            story.append(Paragraph(f"<font color='{color}'><b>{item['jd_skill']}</b></font> ({item['similarity']*100:.1f}% similarity)", body_style))
        story.append(Spacer(1, 8))

    section_list("Matched Skills", bundle.matched_skills, "#2c9e3f")
    section_list("Partial Matches", bundle.partial_matches, "#f39c12")
    section_list("Top Missing Skills", bundle.missing_skills, "#e74c3c")

    if bundle.learning_path:
        story.append(Paragraph("Learning Path Recommendations", section_style))
        for item in bundle.learning_path[:5]:
            story.append(Paragraph(f"<b>{item['skill']}</b> - {item['priority']} priority - {item['estimated_time']}", body_style))
            story.append(Paragraph(item["recommended_action"], body_style))
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()



def create_report_preview_html(bundle: AnalysisBundle) -> str:
    def render_list(items: List[Dict[str, object]], color: str) -> str:
        if not items:
            return "<p>No items found.</p>"
        html_items = []
        for item in items:
            html_items.append(f"<li style='margin-bottom:6px; color:{color};'><b>{item['jd_skill']}</b> ({item['similarity']*100:.1f}% similarity)</li>")
        return "<ul>" + "".join(html_items) + "</ul>"

    return f"""
    <div class="report-shell">
        <h1>Skill Gap Analysis Report</h1>
        <p>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <h2>Overall Match Score</h2>
        <table>
            <tr><th>Match Percentage</th><td>{bundle.match_score:.1f}%</td></tr>
            <tr><th>ATS Score</th><td>{bundle.ats_score:.1f}%</td></tr>
        </table>
        <h2>Skills Summary</h2>
        <table>
            <tr><th>Category</th><th>Resume Skills</th><th>JD Skills</th><th>Matched</th></tr>
            <tr><td>Technical</td><td>{bundle.category_counts_resume.get('technical', 0)}</td><td>{bundle.category_counts_jd.get('technical', 0)}</td><td>{len(bundle.matched_skills)}</td></tr>
        </table>
        <h2>Matched Skills</h2>
        {render_list(bundle.matched_skills, '#2c9e3f')}
        <h2>Partial Matches</h2>
        {render_list(bundle.partial_matches, '#f39c12')}
        <h2>Top Missing Skills</h2>
        {render_list(bundle.missing_skills, '#e74c3c')}
    </div>
    """



def build_report_download(bundle: AnalysisBundle) -> Tuple[bytes, str, str]:
    if REPORTLAB_AVAILABLE:
        pdf_bytes = generate_pdf_report_bytes(bundle)
        if pdf_bytes:
            return pdf_bytes, "application/pdf", "skill_gap_analysis_report.pdf"
    preview = create_report_preview_html(bundle)
    return preview.encode("utf-8"), "text/html", "skill_gap_analysis_report.html"



def render_sidebar() -> None:
    with st.sidebar:
        st.markdown(
            """
            <div class="side-panel">
                <h3>Analysis Settings</h3>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.session_state.analysis_mode = st.selectbox(
            "Extraction Model",
            ["spaCy base model", "AI hybrid model"],
            index=0 if st.session_state.analysis_mode == "spaCy base model" else 1,
        )
        st.session_state.confidence_threshold = st.slider("Confidence Threshold", 0.30, 0.95, float(st.session_state.confidence_threshold), 0.01)
        st.session_state.similarity_threshold = st.slider("Similarity Threshold", 0.30, 0.95, float(st.session_state.similarity_threshold), 0.01)
        st.session_state.embedding_model = st.selectbox(
            "Embedding Model",
            ["all-MiniLM-L6-v2", "paraphrase-MiniLM-L6-v2"],
            index=0 if st.session_state.embedding_model == "all-MiniLM-L6-v2" else 1,
        )

        st.markdown("<hr style='border-color: rgba(255,255,255,0.08);'>", unsafe_allow_html=True)
        st.markdown("<div class='side-panel'><h3>Export Options</h3></div>", unsafe_allow_html=True)
        st.session_state.export_format = st.selectbox("Format", ["PDF", "HTML", "CSV", "JSON"], index=["PDF", "HTML", "CSV", "JSON"].index(st.session_state.export_format))

        if st.button("Generate Report", use_container_width=True):
            bundle = st.session_state.get("bundle")
            if bundle is None:
                st.warning("Process documents first to generate a report.")
            else:
                report_bytes, mime, file_name = build_report_download(bundle)
                st.session_state.last_report_bytes = report_bytes
                st.session_state.last_report_name = file_name
                st.success("Report ready for download.")

        st.markdown("<hr style='border-color: rgba(255,255,255,0.08);'>", unsafe_allow_html=True)
        if st.button("Start New Analysis", use_container_width=True):
            reset_analysis()
            st.session_state.nav_target = "upload"
            st.rerun()

        st.markdown("<hr style='border-color: rgba(255,255,255,0.08);'>", unsafe_allow_html=True)
        st.markdown(
            """
            <div class="small-note">
                AI Skill Gap Analyzer Pro combines document parsing, skill extraction,
                semantic comparison, ATS scoring, learning guidance, and visual analytics.
            </div>
            """,
            unsafe_allow_html=True,
        )



def render_hero() -> None:
    st.markdown(
        """
        <div class="hero">
            <div class="hero-badge">🧠 AI-Powered Career Intelligence</div>
            <h1>AI Skill Gap Analyzer Pro</h1>
            <p>Transform your career with intelligent skill analysis</p>
        </div>
        """,
        unsafe_allow_html=True,
    )



def render_stepper(current_index: int) -> None:
    steps = [
        ("1", "Upload", "📄"),
        ("2", "Extract", "🔎"),
        ("3", "Analyze", "📊"),
        ("4", "Gaps", "🎯"),
        ("5", "Visualize", "📈"),
        ("6", "ATS", "🤖"),
    ]

    step_html = []
    for index, (number, label, icon) in enumerate(steps):
        if index < current_index:
            cls = "done"
        elif index == current_index:
            cls = "active"
        else:
            cls = ""
        step_html.append(
            f"<div class='step-card'><div class='step-circle {cls}'>{number}</div><div class='step-label'>{label}</div><span class='step-icon'>{icon}</span></div>"
        )

    st.markdown(
        f"<div class='step-shell'><div class='step-track'></div><div class='stepper'>{''.join(step_html)}</div></div>",
        unsafe_allow_html=True,
    )

    button_cols = st.columns(6)
    for idx, (number, label, _) in enumerate(steps):
        with button_cols[idx]:
            if st.button(f"Go to {label}", key=f"jump_{number}", use_container_width=True):
                st.session_state.nav_target = label.lower()
                st.toast(f"Jumping to {label}")



def render_upload_area() -> Tuple[Optional[object], Optional[object], Optional[str], Optional[str]]:
    st.markdown("<div class='section-title'>Step 1: Upload Your Documents</div>", unsafe_allow_html=True)
    st.markdown("<div class='section-subtitle'>Upload a resume and a job description, or paste text directly for either field.</div>", unsafe_allow_html=True)

    col_resume, col_jd = st.columns(2)
    resume_file = None
    jd_file = None
    resume_text = ""
    jd_text = ""

    with col_resume:
        st.markdown("<div class='upload-card'><div class='card-title'>Resume</div><div class='upload-zone'>", unsafe_allow_html=True)
        resume_file = st.file_uploader(
            "Drag and drop file here",
            type=["pdf", "docx", "txt"],
            key="resume_uploader",
            label_visibility="collapsed",
        )
        resume_paste = st.checkbox("Or paste text directly", key="resume_paste")
        if resume_paste:
            resume_text = st.text_area("Resume text", height=170, key="resume_text_input")
        st.markdown("</div></div>", unsafe_allow_html=True)

    with col_jd:
        st.markdown("<div class='upload-card'><div class='card-title'>Job Description</div><div class='upload-zone'>", unsafe_allow_html=True)
        jd_file = st.file_uploader(
            "Drag and drop file here",
            type=["pdf", "docx", "txt"],
            key="jd_uploader",
            label_visibility="collapsed",
        )
        jd_paste = st.checkbox("Or paste text directly", key="jd_paste")
        if jd_paste:
            jd_text = st.text_area("Job description text", height=170, key="jd_text_input")
        st.markdown("</div></div>", unsafe_allow_html=True)

    return resume_file, jd_file, resume_text, jd_text



def process_documents(resume_file, jd_file, resume_text: str, jd_text: str) -> Optional[AnalysisBundle]:
    if resume_file is None and not resume_text.strip():
        st.warning("Add a resume file or paste resume text.")
        return None
    if jd_file is None and not jd_text.strip():
        st.warning("Add a job description file or paste job description text.")
        return None

    resume_input = None
    jd_input = None

    if resume_file is not None:
        resume_input = {
            "name": resume_file.name,
            "bytes": resume_file.getvalue(),
            "kind": "resume",
        }
    elif resume_text.strip():
        resume_input = {
            "name": "pasted_resume.txt",
            "bytes": resume_text.encode("utf-8"),
            "kind": "resume",
        }

    if jd_file is not None:
        jd_input = {
            "name": jd_file.name,
            "bytes": jd_file.getvalue(),
            "kind": "job_description",
        }
    elif jd_text.strip():
        jd_input = {
            "name": "pasted_job_description.txt",
            "bytes": jd_text.encode("utf-8"),
            "kind": "job_description",
        }

    if resume_input is None or jd_input is None:
        return None

    with st.spinner("Processing documents..."):
        resume_doc = parse_and_clean_document(resume_input["name"], resume_input["bytes"], "resume")
        jd_doc = parse_and_clean_document(jd_input["name"], jd_input["bytes"], "job_description")

        resume_skills, advanced_resume_skills, resume_confidence = extract_skills_hybrid(resume_doc.cleaned_text, st.session_state.analysis_mode)
        jd_skills, advanced_jd_skills, jd_confidence = extract_skills_hybrid(jd_doc.cleaned_text, st.session_state.analysis_mode)

        if not resume_skills and advanced_resume_skills:
            resume_skills = advanced_resume_skills
        if not jd_skills and advanced_jd_skills:
            jd_skills = advanced_jd_skills

        resume_doc.extracted_skills = resume_skills
        resume_doc.advanced_skills = advanced_resume_skills
        resume_doc.skill_confidence = resume_confidence
        jd_doc.extracted_skills = jd_skills
        jd_doc.advanced_skills = advanced_jd_skills
        jd_doc.skill_confidence = jd_confidence

        category_resume = build_category_counts(resume_skills)
        category_jd = build_category_counts(jd_skills)

        analysis_result, matched, partial, missing, similarity_matrix, match_score = run_semantic_gap_analysis(
            resume_skills,
            jd_skills,
            st.session_state.similarity_threshold,
        )

        matched_sorted = sorted(matched, key=lambda item: item.get("similarity", 0), reverse=True)
        partial_sorted = sorted(partial, key=lambda item: item.get("similarity", 0), reverse=True)
        missing_sorted = sorted(missing, key=lambda item: item.get("similarity", 0))

        ats_score, ats_factors = calculate_ats_factors(
            resume_doc.cleaned_text,
            jd_doc.cleaned_text,
            resume_skills,
            jd_skills,
            None,
        )

        learning_path = build_learning_path(missing_sorted)

        bundle = AnalysisBundle(
            resume=resume_doc,
            job_description=jd_doc,
            all_resume_skills=resume_skills,
            all_jd_skills=jd_skills,
            analysis_result=analysis_result,
            match_score=round(match_score * 100, 1),
            ats_score=ats_score,
            ats_factors=ats_factors,
            matched_skills=matched_sorted,
            partial_matches=partial_sorted,
            missing_skills=missing_sorted,
            learning_path=learning_path,
            category_counts_resume=category_resume,
            category_counts_jd=category_jd,
            similarity_matrix=similarity_matrix,
            skill_pair_map=matched_sorted + partial_sorted + missing_sorted,
        )
        bundle.ats_score, bundle.ats_factors = calculate_ats_factors(
            resume_doc.cleaned_text,
            jd_doc.cleaned_text,
            resume_skills,
            jd_skills,
            bundle,
        )

    st.session_state.bundle = bundle
    st.session_state.processed_docs = [resume_doc, jd_doc]
    st.session_state.nav_target = "extract"
    return bundle



def render_processing_status(bundle: AnalysisBundle) -> None:
    st.markdown("<div class='section-title'>File Processing Status</div>", unsafe_allow_html=True)
    for doc in [bundle.resume, bundle.job_description]:
        icon = "📄" if doc.doc_type == "resume" else "💼"
        status_class = "good"
        st.markdown(
            f"""
            <div class="status-card {status_class}">
                <div style="display:flex; align-items:center; gap: 12px;">
                    <div style="font-size: 1.25rem;">{icon}</div>
                    <div>
                        <div style="color:white; font-weight:700; font-size:1rem;">{doc.name}</div>
                        <div style="color:rgba(255,255,255,0.72); font-size:0.88rem;">
                            Type: {doc.name.split('.')[-1].upper()} | Size: {doc.file_size_kb} KB | Pages: {doc.pages} | Words: {doc.word_count} | Skills: {len(doc.extracted_skills)}
                        </div>
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("<div class='section-title'>Normalization Summary</div>", unsafe_allow_html=True)
    left, right = st.columns(2)
    with left:
        st.markdown(
            f"""
            <div class="status-card">
                <div style="color:white; font-weight:700; margin-bottom: 8px;">Resume</div>
                <div class="small-note">Original: {bundle.resume.original_chars} chars → Cleaned: {bundle.resume.cleaned_chars} chars</div>
                <div class="small-note">Removed: {bundle.resume.removed_lines} lines, {bundle.resume.removed_chars} characters</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with right:
        st.markdown(
            f"""
            <div class="status-card">
                <div style="color:white; font-weight:700; margin-bottom: 8px;">Job Description</div>
                <div class="small-note">Original: {bundle.job_description.original_chars} chars → Cleaned: {bundle.job_description.cleaned_chars} chars</div>
                <div class="small-note">Removed: {bundle.job_description.removed_lines} lines, {bundle.job_description.removed_chars} characters</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("<div class='section-title'>Extracted Skills</div>", unsafe_allow_html=True)
    metric_cols = st.columns(3)
    metric_cols[0].markdown(
        f"<div class='metric-box'><div class='metric-big'>{len(bundle.all_resume_skills)}</div><div class='metric-label'>Resume Skills</div></div>",
        unsafe_allow_html=True,
    )
    metric_cols[1].markdown(
        f"<div class='metric-box'><div class='metric-big'>{len(bundle.all_jd_skills)}</div><div class='metric-label'>Job Description Skills</div></div>",
        unsafe_allow_html=True,
    )
    metric_cols[2].markdown(
        f"<div class='metric-box'><div class='metric-big'>{bundle.match_score:.1f}%</div><div class='metric-label'>Match Score</div></div>",
        unsafe_allow_html=True,
    )

    st.markdown("<div style='margin-top:0.9rem;'></div>", unsafe_allow_html=True)
    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("<div class='card-title'>Resume Skills</div>", unsafe_allow_html=True)
        st.markdown(create_tag_html(bundle.all_resume_skills, "technical"), unsafe_allow_html=True)
    with col_b:
        st.markdown("<div class='card-title'>Job Description Skills</div>", unsafe_allow_html=True)
        st.markdown(create_tag_html(bundle.all_jd_skills, "technical"), unsafe_allow_html=True)



def render_gap_analysis(bundle: AnalysisBundle) -> None:
    st.markdown("<div class='section-title'>Gap Analysis Results</div>", unsafe_allow_html=True)
    metric_cols = st.columns(4)
    metric_cols[0].markdown(
        f"<div class='metric-box'><div class='metric-big'>{bundle.match_score:.1f}%</div><div class='metric-label'>Match Score</div></div>",
        unsafe_allow_html=True,
    )
    metric_cols[1].markdown(
        f"<div class='metric-box'><div class='metric-big'>{len(bundle.matched_skills)}</div><div class='metric-label'>Matched</div></div>",
        unsafe_allow_html=True,
    )
    metric_cols[2].markdown(
        f"<div class='metric-box'><div class='metric-big'>{len(bundle.partial_matches)}</div><div class='metric-label'>Partial</div></div>",
        unsafe_allow_html=True,
    )
    metric_cols[3].markdown(
        f"<div class='metric-box'><div class='metric-big'>{len(bundle.missing_skills)}</div><div class='metric-label'>Missing</div></div>",
        unsafe_allow_html=True,
    )

    col_left, col_mid, col_right = st.columns(3)
    with col_left:
        st.markdown("<div class='card-title'>Matched Skills</div>", unsafe_allow_html=True)
        st.markdown(create_tag_html([item['jd_skill'] for item in bundle.matched_skills], "technical"), unsafe_allow_html=True)
    with col_mid:
        st.markdown("<div class='card-title'>Partial Matches</div>", unsafe_allow_html=True)
        st.markdown(create_tag_html([f"{item['jd_skill']} ({item['similarity']*100:.1f}%)" for item in bundle.partial_matches], "soft"), unsafe_allow_html=True)
    with col_right:
        st.markdown("<div class='card-title'>Missing Skills</div>", unsafe_allow_html=True)
        st.markdown(create_tag_html([f"{item['jd_skill']} ({item['similarity']*100:.1f}%)" for item in bundle.missing_skills], "certifications"), unsafe_allow_html=True)



def render_learning_path(bundle: AnalysisBundle) -> None:
    st.markdown("<div class='section-title'>Learning Path</div>", unsafe_allow_html=True)
    st.markdown("<div class='section-subtitle'>Recommended next steps based on your missing skills and priority order.</div>", unsafe_allow_html=True)

    if not bundle.learning_path:
        st.info("No major gaps detected. The resume already covers the core requirements well.")
        return

    nav_cols = st.columns([1, 1, 1, 1, 1, 1])
    nav_steps = ["Upload", "Extract", "Analyze", "Gaps", "Visualize", "ATS"]
    for idx, label in enumerate(nav_steps):
        with nav_cols[idx]:
            if st.button(f"Go to {label}", key=f"learning_nav_{label}", use_container_width=True):
                st.session_state.nav_target = label.lower()

    for index, item in enumerate(bundle.learning_path, start=1):
        st.markdown(
            f"""
            <div class="status-card">
                <div style="border-left: 4px solid #f7c948; padding-left: 16px;">
                    <div style="color:white; font-size:1.05rem; font-weight:800; margin-bottom:6px;">{index}. {item['skill']} ({item['priority']} Priority)</div>
                    <div class="small-note">Importance: {item['importance']}</div>
                    <div class="small-note">Estimated Time: {item['estimated_time']}</div>
                    <div class="small-note" style="margin-top:8px;">Recommended Action: {item['recommended_action']}</div>
                    <div style="margin-top:10px; color:white; font-weight:700;">Resources:</div>
                    <ul style="margin-top:6px; color:rgba(255,255,255,0.85);">
                        {''.join(f'<li>{resource}</li>' for resource in item['resources'])}
                    </ul>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )



def render_visualization(bundle: AnalysisBundle) -> None:
    st.markdown("<div class='section-title'>Visualisation</div>", unsafe_allow_html=True)
    st.markdown("<div class='section-subtitle'>Switch between tag cloud, radar chart, heatmap, and distribution to inspect the skill profile.</div>", unsafe_allow_html=True)

    tabs = st.tabs(["Tag Cloud", "Radar Chart", "Heatmap", "Distribution"])

    with tabs[0]:
        st.markdown("<div class='card-title'>Skills Tag Cloud</div>", unsafe_allow_html=True)
        cloud = generate_wordcloud_image(bundle.all_resume_skills + bundle.all_jd_skills)
        if cloud is None:
            st.info("No data available for the tag cloud.")
        else:
            fig, ax = plt.subplots(figsize=(14, 7), facecolor="white")
            ax.imshow(cloud, interpolation="bilinear")
            ax.axis("off")
            st.pyplot(fig, use_container_width=True)
            plt.close(fig)

    with tabs[1]:
        st.markdown("<div class='card-title'>Skill Category Comparison</div>", unsafe_allow_html=True)
        st.plotly_chart(generate_radar_chart(bundle.category_counts_resume, bundle.category_counts_jd), use_container_width=True)

    with tabs[2]:
        st.markdown("<div class='card-title'>Skill Similarity Heatmap</div>", unsafe_allow_html=True)
        st.plotly_chart(generate_similarity_heatmap(bundle.all_resume_skills, bundle.all_jd_skills, bundle.similarity_matrix), use_container_width=True)

    with tabs[3]:
        st.markdown("<div class='card-title'>Skill Distribution</div>", unsafe_allow_html=True)
        st.plotly_chart(generate_distribution_chart(bundle.category_counts_resume), use_container_width=True)



def render_ats(bundle: AnalysisBundle) -> None:
    st.markdown("<div class='section-title'>ATS Score</div>", unsafe_allow_html=True)
    score_class = score_badge_class(bundle.ats_score)
    status_text = "Excellent" if bundle.ats_score >= 80 else "Good" if bundle.ats_score >= 60 else "Needs Improvement" if bundle.ats_score >= 40 else "Poor Match"

    st.markdown(
        f"""
        <div class="score-card">
            <div class="score-circle {score_class}">{bundle.ats_score:.0f}%</div>
            <div style="color:white; font-weight:800; font-size:1.2rem; margin-bottom:0.25rem;">{status_text}</div>
            <div class="small-note">Your resume is analyzed against keyword coverage, formatting, structure, readability, and relevance.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("<div class='section-title' style='margin-top:1rem;'>ATS Optimization Factors</div>", unsafe_allow_html=True)
    for key, label in [
        ("keyword_match", "Keyword Match"),
        ("formatting", "Formatting"),
        ("section_completeness", "Section Completeness"),
        ("readability", "Readability"),
        ("experience_relevance", "Experience Relevance"),
    ]:
        factor = bundle.ats_factors[key]
        score = float(factor["score"])
        category = str(factor["category"])
        color = CATEGORY_COLORS["technical"] if category == "excellent" else CATEGORY_COLORS["tools"] if category == "average" else CATEGORY_COLORS["certifications"]
        st.markdown(
            f"""
            <div class="ats-shell">
                <div class="ats-header">
                    <div class="ats-title">{label}</div>
                    <div class="chip {'good' if category == 'excellent' else 'warn' if category == 'average' else 'bad'}">{score:.0f}%</div>
                </div>
                <div class="ats-bar"><div style="width:{score}%; background:{color};"></div></div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        with st.expander(f"Recommendations for {label}"):
            st.write(factor["recommendation"])

    if bundle.missing_skills:
        st.markdown("<div class='section-title' style='margin-top:1rem;'>Missing Keywords</div>", unsafe_allow_html=True)
        st.markdown(
            "<div class='small-note'>Consider adding these keywords from the job description to your resume:</div>",
            unsafe_allow_html=True,
        )
        st.markdown(create_tag_html([item["jd_skill"] for item in bundle.missing_skills[:8]], "soft"), unsafe_allow_html=True)



def render_report_section(bundle: AnalysisBundle) -> None:
    st.markdown("<div class='section-title'>Report</div>", unsafe_allow_html=True)
    st.markdown("<div class='section-subtitle'>Generate a downloadable report based on the current analysis output.</div>", unsafe_allow_html=True)

    report_bytes, mime, file_name = build_report_download(bundle)
    if bundle.ats_score and bundle.match_score:
        st.download_button("Download Report", data=report_bytes, file_name=file_name, mime=mime, use_container_width=True)

    with st.expander("Preview report"):
        if mime == "text/html":
            st.components.v1.html(report_bytes.decode("utf-8", errors="ignore"), height=900, scrolling=True)
        else:
            st.write("PDF report is ready for download.")
        st.markdown("### Summary")
        st.write(
            {
                "match_score": bundle.match_score,
                "ats_score": bundle.ats_score,
                "matched": len(bundle.matched_skills),
                "partial": len(bundle.partial_matches),
                "missing": len(bundle.missing_skills),
                "learning_items": len(bundle.learning_path),
            }
        )



def main() -> None:
    ensure_state()
    load_custom_css()
    render_sidebar()
    render_hero()

    stage_map = {
        "upload": 0,
        "extract": 1,
        "analyze": 2,
        "gaps": 3,
        "visualize": 4,
        "ats": 5,
    }
    current_stage = stage_map.get(st.session_state.nav_target, 0)
    if st.session_state.bundle is not None:
        current_stage = max(current_stage, 1)

    render_stepper(current_stage)

    if st.button("Load Demo Case", use_container_width=False):
        demo_inputs = get_sample_inputs_from_workspace()
        if not demo_inputs:
            st.error("Demo files were not found in the workspace.")
        else:
            resume_input, jd_input = demo_inputs
            with st.spinner("Loading demo analysis..."):
                resume_doc = parse_and_clean_document(resume_input["name"], resume_input["bytes"], "resume")
                jd_doc = parse_and_clean_document(jd_input["name"], jd_input["bytes"], "job_description")
                resume_skills, advanced_resume_skills, resume_confidence = extract_skills_hybrid(resume_doc.cleaned_text, st.session_state.analysis_mode)
                jd_skills, advanced_jd_skills, jd_confidence = extract_skills_hybrid(jd_doc.cleaned_text, st.session_state.analysis_mode)
                resume_doc.extracted_skills = resume_skills
                resume_doc.advanced_skills = advanced_resume_skills
                resume_doc.skill_confidence = resume_confidence
                jd_doc.extracted_skills = jd_skills
                jd_doc.advanced_skills = advanced_jd_skills
                jd_doc.skill_confidence = jd_confidence
                category_resume = build_category_counts(resume_skills)
                category_jd = build_category_counts(jd_skills)
                analysis_result, matched, partial, missing, similarity_matrix, match_score = run_semantic_gap_analysis(resume_skills, jd_skills, st.session_state.similarity_threshold)
                matched_sorted = sorted(matched, key=lambda item: item.get("similarity", 0), reverse=True)
                partial_sorted = sorted(partial, key=lambda item: item.get("similarity", 0), reverse=True)
                missing_sorted = sorted(missing, key=lambda item: item.get("similarity", 0))
                temp_bundle = AnalysisBundle(
                    resume=resume_doc,
                    job_description=jd_doc,
                    all_resume_skills=resume_skills,
                    all_jd_skills=jd_skills,
                    analysis_result=analysis_result,
                    match_score=round(match_score * 100, 1),
                    ats_score=0.0,
                    ats_factors={},
                    matched_skills=matched_sorted,
                    partial_matches=partial_sorted,
                    missing_skills=missing_sorted,
                    learning_path=build_learning_path(missing_sorted),
                    category_counts_resume=category_resume,
                    category_counts_jd=category_jd,
                    similarity_matrix=similarity_matrix,
                    skill_pair_map=matched_sorted + partial_sorted + missing_sorted,
                )
                temp_bundle.ats_score, temp_bundle.ats_factors = calculate_ats_factors(resume_doc.cleaned_text, jd_doc.cleaned_text, resume_skills, jd_skills, temp_bundle)
                st.session_state.bundle = temp_bundle
                st.session_state.processed_docs = [resume_doc, jd_doc]
                st.session_state.demo_loaded = True
                st.session_state.nav_target = "extract"
                st.success("Demo analysis loaded.")
                st.rerun()

    resume_file, jd_file, resume_text, jd_text = render_upload_area()

    process_col, filler = st.columns([1, 3])
    with process_col:
        if st.button("Process Documents", use_container_width=True):
            bundle = process_documents(resume_file, jd_file, resume_text, jd_text)
            if bundle is not None:
                st.success("Documents processed successfully.")
                st.rerun()
    with filler:
        st.markdown(
            "<div class='small-note' style='padding-top:0.5rem;'>Process both documents to unlock analysis, gap detection, visualization, ATS scoring, and report export.</div>",
            unsafe_allow_html=True,
        )

    bundle = st.session_state.get("bundle")
    if bundle is None:
        st.info("Upload a resume and job description to begin the full analysis flow, or load the demo case to preview the completed pipeline.")
        return

    st.markdown("<div style='height: 1rem;'></div>", unsafe_allow_html=True)
    render_processing_status(bundle)
    st.markdown("<div style='height: 0.4rem;'></div>", unsafe_allow_html=True)
    render_gap_analysis(bundle)
    st.markdown("<div style='height: 0.4rem;'></div>", unsafe_allow_html=True)
    render_learning_path(bundle)
    st.markdown("<div style='height: 0.4rem;'></div>", unsafe_allow_html=True)
    render_visualization(bundle)
    st.markdown("<div style='height: 0.4rem;'></div>", unsafe_allow_html=True)
    render_ats(bundle)
    st.markdown("<div style='height: 0.4rem;'></div>", unsafe_allow_html=True)
    render_report_section(bundle)

    st.markdown(
        "<div class='footer-note'>AI Skill Gap Analyzer Pro - resume parsing, semantic comparison, ATS scoring, and visual analytics.</div>",
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
